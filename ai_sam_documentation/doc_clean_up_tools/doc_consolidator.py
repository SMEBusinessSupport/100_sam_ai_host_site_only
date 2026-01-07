#!/usr/bin/env python3
"""
Documentation Consolidator Script
Auto-consolidates scattered .md files into organized structure.

Usage:
    python doc_consolidator.py run       # Auto-move all files + generate report
    python doc_consolidator.py --help    # Show help

Reports saved to: clean_up_reports/
"""

import os
import sys
import csv
import re
import shutil
from pathlib import Path
from datetime import datetime
from collections import defaultdict

# Fix Windows console encoding issues
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# Optional: python-docx for Word document conversion
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ============================================================================
# CONFIGURATION
# ============================================================================

CONFIG = {
    # Source folders to scan for .md files
    'source_dirs': [
        # Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_documentation/docs/_archive/unsorted'),
        # Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_documentation/old_docs'),
        # Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam'),
        # Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam/chat_ui_wip'),
        # Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_base'),
        # Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_workflows'),
        # Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_workflows_base'),
        # Path(r'D:/SAMAI-18-SaaS/ai_sam/ai_brain/docs'),
        # Path(r'D:/SAMAI-18-SaaS/ai_sam/ai_sam/documentation'),
        # Path(r'D:/SAMAI-18-SaaS/ai_sam/ai_sam/dev docs'),
        # Path(r'D:/SAMAI-18-SaaS/ai_sam/ai_sam_workflows/documentation'),
        # Path(r'D:/SAMAI-18-SaaS/ai_sam/ai_sam_intelligence/dev docs'),
        # === 12-samai-docs-and-tools FULL SWEEP ===
        # Path(r'D:/SAMAI-18-SaaS/github-repos/12-samai-docs-and-tools'),
        # === ai_sam_introduction - GOLD CONTENT ===
        Path(r'D:/SAMAI-18-SaaS/github-repos/ai_sam_introduction'),  # Current sweep
    ],

    # Destination root (where _README.md files live)
    'dest_root': Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_documentation/docs'),

    # Where cleanup reports are saved
    'reports_dir': Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_documentation/clean_up_reports'),

    # Fallback for unmatched files (very low confidence only)
    'unmatched_dest': '_unsorted',

    # Asset folder name (for images, diagrams, etc.)
    'assets_folder': '_assets',

    # Asset file extensions (images, diagrams)
    'asset_extensions': ['.png', '.jpg', '.jpeg', '.gif', '.svg', '.webp', '.ico', '.bmp'],

    # Script file extensions (convert to .md with code blocks)
    'script_extensions': {
        '.sql': 'sql',
        '.json': 'json',
        '.py': 'python',
        '.js': 'javascript',
        '.xml': 'xml',
        '.css': 'css',
        '.yaml': 'yaml',
        '.yml': 'yaml',
        '.mermaid': 'mermaid',
    },

    # HTML files - MOVE as-is (don't convert to markdown)
    'html_extensions': ['.html', '.htm'],

    # Where to store original script files (within module)
    'scripts_backup_dir': Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_documentation/doc_scripts'),

    # Word document conversion (requires python-docx)
    'docx_enabled': True,

    # Shallow mode: only process files directly in source folders (not subfolders)
    'shallow_mode': False,  # DISABLED for full recursive sweep

    # Old tools directory (for obsolete cleanup scripts, bat files, etc.)
    'old_tools_dir': Path(r'D:/SAMAI-18-SaaS/github-repos/05-samai-core/ai_sam_documentation/cleaning_up_old_script_tools'),

    # Patterns to identify old tool files (NOT module files - things like diagnose_brain.py, cleanup.bat)
    # These are files that are clearly tools/utilities rather than module code
    'old_tool_patterns': [
        r'.*diagnose.*\.py$',      # Diagnostic tools
        r'.*cleanup.*\.py$',       # Cleanup tools
        r'.*debug.*\.py$',         # Debug tools
        r'.*fix.*\.py$',           # Fix scripts
        r'.*migrate.*\.py$',       # Migration scripts
        r'.*audit.*\.py$',         # Audit scripts
        r'.*check.*\.py$',         # Check scripts
        r'.*test_.*\.py$',         # Test scripts (standalone)
        r'.*_tool\.py$',           # Anything ending in _tool.py
        r'.*_script\.py$',         # Anything ending in _script.py
        r'.*\.bat$',               # All batch files
        r'.*\.cmd$',               # All command files
        r'.*\.sh$',                # All shell scripts
    ],

    # Files/patterns to skip entirely
    'skip_patterns': [
        r'^_README\.md$',      # Skip category README files
        r'^README\.md$',       # Skip module README files
        r'^\..*',              # Skip hidden files
        r'__pycache__',        # Skip cache
        r'^__init__\.py$',     # Skip Odoo module init
        r'^__manifest__\.py$', # Skip Odoo module manifest
        r'^hooks\.py$',        # Skip Odoo module hooks
        r'\.pyc$',             # Skip compiled Python
    ],

    # Confidence thresholds
    'high_confidence': 0.7,    # 70%+ = high confidence (auto-move, no flag)
    'medium_confidence': 0.3,  # 30-70% = medium confidence (auto-move, flag for review)
    'very_low_confidence': 0.15,  # Below 15% = very low (move to _unsorted/, flag)
                               # 15-30% = low confidence (still move to matched folder, flag)
}


# ============================================================================
# CATEGORY PARSER
# ============================================================================

class CategoryParser:
    """Parses _README.md files to extract matching keywords."""

    def __init__(self, dest_root: Path):
        self.dest_root = dest_root
        self.categories = {}
        self._parse_all_readmes()

    def _parse_all_readmes(self):
        """Find and parse all _README.md files."""
        for readme_path in self.dest_root.rglob('_README.md'):
            rel_folder = readme_path.parent.relative_to(self.dest_root)
            self.categories[str(rel_folder)] = self._parse_readme(readme_path)

    def _parse_readme(self, path: Path) -> dict:
        """Extract keywords from a _README.md file."""
        content = path.read_text(encoding='utf-8')

        # Extract sections
        purpose = self._extract_section(content, 'Purpose')
        criteria = self._extract_section(content, 'Criteria')
        examples = self._extract_section(content, 'Examples')
        excludes = self._extract_section(content, 'Does NOT Include')
        subfolders = self._extract_section(content, 'Subfolders')

        # Build keyword lists
        include_keywords = self._extract_keywords(purpose + criteria + examples)
        exclude_keywords = self._extract_keywords(excludes)
        subfolder_names = self._extract_subfolder_names(subfolders)

        return {
            'path': path.parent,
            'include_keywords': include_keywords,
            'exclude_keywords': exclude_keywords,
            'subfolders': subfolder_names,
            'raw_purpose': purpose,
        }

    def _extract_section(self, content: str, header: str) -> str:
        """Extract content under a ## header."""
        pattern = rf'## {header}\s*\n(.*?)(?=\n## |\Z)'
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        return match.group(1).strip() if match else ''

    def _extract_keywords(self, text: str) -> list:
        """Extract meaningful keywords from text."""
        # Remove markdown formatting
        text = re.sub(r'[*_`#\[\]()]', ' ', text)
        text = re.sub(r'-\s+', ' ', text)

        # Extract words (3+ chars, lowercase)
        words = re.findall(r'\b[a-zA-Z_]{3,}\b', text.lower())

        # Filter common words
        stopwords = {'the', 'and', 'for', 'that', 'this', 'with', 'are', 'from',
                     'have', 'has', 'how', 'what', 'why', 'when', 'where', 'which',
                     'can', 'will', 'should', 'would', 'could', 'does', 'not',
                     'into', 'about', 'than', 'them', 'then', 'each', 'other'}

        return [w for w in words if w not in stopwords]

    def _extract_subfolder_names(self, text: str) -> list:
        """Extract subfolder names from Subfolders section."""
        # Match patterns like: `folder/` or - `folder/`
        matches = re.findall(r'`([a-zA-Z_]+)/?`', text)
        return matches


# ============================================================================
# FILE SCORER
# ============================================================================

class FileScorer:
    """Scores source files against destination categories."""

    def __init__(self, categories: dict):
        self.categories = categories

    def score_file(self, file_path: Path) -> list:
        """
        Score a file against all categories.
        Returns list of (category, score, reason) tuples, sorted by score.
        """
        # Get file info
        filename = file_path.name.lower()
        filename_keywords = self._extract_file_keywords(filename)

        # Read first 1000 chars of content
        try:
            content = file_path.read_text(encoding='utf-8')[:1000].lower()
        except:
            content = ''

        content_keywords = self._extract_content_keywords(content)

        # Get source folder context
        source_folder = file_path.parent.name.lower()

        results = []

        for cat_name, cat_info in self.categories.items():
            score, reasons = self._calculate_score(
                filename_keywords,
                content_keywords,
                source_folder,
                cat_info
            )
            results.append((cat_name, score, reasons))

        # Sort by score descending
        results.sort(key=lambda x: x[1], reverse=True)
        return results

    def _extract_file_keywords(self, filename: str) -> list:
        """Extract keywords from filename."""
        # Remove extension, split on separators
        name = re.sub(r'\.[^.]+$', '', filename)
        words = re.split(r'[_\-\s]+', name)
        return [w.lower() for w in words if len(w) >= 2]

    def _extract_content_keywords(self, content: str) -> list:
        """Extract keywords from file content."""
        # Get first heading
        heading_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        heading = heading_match.group(1).lower() if heading_match else ''

        # Extract words
        words = re.findall(r'\b[a-zA-Z_]{3,}\b', content)
        words = [w.lower() for w in words[:100]]  # First 100 words

        return words + heading.split()

    def _calculate_score(self, filename_kw, content_kw, source_folder, cat_info) -> tuple:
        """Calculate match score and reasons."""
        score = 0.0
        reasons = []

        include_kw = set(cat_info['include_keywords'])
        exclude_kw = set(cat_info['exclude_keywords'])

        # Filename matches (weighted heavily)
        filename_matches = set(filename_kw) & include_kw
        if filename_matches:
            score += len(filename_matches) * 0.15
            reasons.append(f"filename:{','.join(filename_matches)}")

        # Content matches
        content_matches = set(content_kw) & include_kw
        if content_matches:
            score += min(len(content_matches) * 0.05, 0.3)
            top_matches = list(content_matches)[:3]
            reasons.append(f"content:{','.join(top_matches)}")

        # Source folder context
        if source_folder in str(cat_info['path']).lower():
            score += 0.2
            reasons.append(f"folder_match:{source_folder}")

        # Subfolder matches
        for subfolder in cat_info.get('subfolders', []):
            if subfolder.lower() in filename_kw or subfolder.lower() in content_kw:
                score += 0.15
                reasons.append(f"subfolder:{subfolder}")

        # Exclusion penalty
        exclude_matches = (set(filename_kw) | set(content_kw)) & exclude_kw
        if exclude_matches:
            score -= len(exclude_matches) * 0.1
            reasons.append(f"excluded:{','.join(exclude_matches)}")

        # Cap score at 1.0
        score = max(0, min(1.0, score))

        return score, reasons


# ============================================================================
# CONSOLIDATOR (Auto-Move + Report)
# ============================================================================

class Consolidator:
    """Auto-moves files and generates cleanup report."""

    def __init__(self, config: dict):
        self.config = config
        self.parser = CategoryParser(config['dest_root'])
        self.scorer = FileScorer(self.parser.categories)
        self.move_log = []
        self.asset_log = []
        self.script_log = []
        self.docx_log = []
        self.old_tools_log = []
        self.html_log = []
        self.stats = {'moved': 0, 'skipped': 0, 'errors': 0, 'high': 0, 'medium': 0, 'low': 0, 'very_low': 0}
        self.asset_stats = {'moved': 0, 'errors': 0}
        self.script_stats = {'converted': 0, 'errors': 0}
        self.docx_stats = {'converted': 0, 'skipped': 0, 'errors': 0}
        self.old_tools_stats = {'moved': 0, 'errors': 0}
        self.html_stats = {'moved': 0, 'errors': 0}
        self.shallow_mode = self.config.get('shallow_mode', False)

    def _get_files(self, source_dir: Path, pattern: str = '*') -> list:
        """Get files from source directory. Respects shallow_mode setting."""
        if self.shallow_mode:
            # Only files directly in the folder
            return list(source_dir.glob(pattern))
        else:
            # Recursive
            return list(source_dir.rglob(pattern))

    def run(self):
        """Scan all source files, auto-move, and generate report."""
        print("\n=== Documentation Consolidator: AUTO-MOVE Mode ===\n")
        print("Philosophy: Move everything, report what we did, human cleans up after.")
        if self.shallow_mode:
            print("Mode: SHALLOW (only files directly in folder, no subfolders)\n")
        else:
            print("Mode: RECURSIVE (includes subfolders)\n")

        for source_dir in self.config['source_dirs']:
            if not source_dir.exists():
                print(f"  [SKIP] Source not found: {source_dir}")
                continue

            print(f"  Processing .md files: {source_dir}")
            self._process_directory(source_dir)

        # Process asset files (images, diagrams)
        print("\n  Processing asset files (images, diagrams)...")
        self._process_assets()

        # Process script files (convert to .md)
        print("\n  Processing script files (sql, json, etc.)...")
        self._process_scripts()

        # Process Word documents (convert to .md)
        print("\n  Processing Word documents (.docx)...")
        self._process_docx()

        # Process HTML files (MOVE as-is, don't convert)
        print("\n  Processing HTML files (move as-is)...")
        self._process_html_files()

        # Process old tool files (move to cleaning_up_old_script_tools/)
        print("\n  Processing old tool files (.py, .bat, .sh)...")
        self._process_old_tools()

        # Check for remaining content
        print("\n  Checking for remaining content...")
        remaining_files = self._find_remaining_content()

        # Clean up empty folders in source directories
        print("\n  Cleaning up empty folders...")
        empty_removed = self._cleanup_empty_folders()

        # Generate report
        report_path = self._write_report()

        # Print summary
        self._print_summary(report_path, empty_removed, remaining_files)

        return report_path

    def _process_assets(self):
        """Process and move asset files (images, diagrams) to _assets/ folders."""
        asset_extensions = set(self.config.get('asset_extensions', []))
        assets_folder = self.config.get('assets_folder', '_assets')

        for source_dir in self.config['source_dirs']:
            if not source_dir.exists():
                continue

            for item in self._get_files(source_dir):
                if not item.is_file():
                    continue

                # Check if it's an asset file
                if item.suffix.lower() not in asset_extensions:
                    continue

                # Skip files matching skip patterns
                if self._should_skip(item):
                    continue

                # Try to find related category based on:
                # 1. Source folder name matching a category
                # 2. Filename matching category keywords
                # 3. Default to root _assets/
                dest_category = self._find_asset_category(item)

                # Build destination path
                if dest_category:
                    dest_path = self.config['dest_root'] / dest_category / assets_folder / item.name
                else:
                    dest_path = self.config['dest_root'] / assets_folder / item.name

                # Execute move
                status, final_dest = self._move_file(item, dest_path)

                # Log the asset move
                self.asset_log.append({
                    'source_path': str(item),
                    'filename': item.name,
                    'file_type': item.suffix,
                    'destination': str(final_dest) if final_dest else str(dest_path),
                    'category': dest_category or 'root',
                    'status': status,
                })

                if status == 'SUCCESS':
                    self.asset_stats['moved'] += 1
                    print(f"    [ASSET] {item.name} -> {dest_category or 'root'}/_assets/")
                else:
                    self.asset_stats['errors'] += 1

    def _process_scripts(self):
        """Convert script files to .md and backup originals to doc_scripts/."""
        script_extensions = self.config.get('script_extensions', {})
        scripts_backup = self.config.get('scripts_backup_dir')

        if not scripts_backup:
            return

        for source_dir in self.config['source_dirs']:
            if not source_dir.exists():
                continue

            for item in self._get_files(source_dir):
                if not item.is_file():
                    continue

                # Check if it's a script file
                ext_lower = item.suffix.lower()
                if ext_lower not in script_extensions:
                    continue

                # Skip temp files (like ~$*.docx)
                if item.name.startswith('~$'):
                    continue

                # Skip files matching skip patterns (module files, etc.)
                if self._should_skip(item):
                    continue

                try:
                    # Read original content
                    content = item.read_text(encoding='utf-8')
                    lang = script_extensions[ext_lower]

                    # Create markdown version
                    md_content = self._create_script_markdown(item.name, content, lang)

                    # Find category for the .md file
                    dest_category = self._find_script_category(item, content)

                    # Determine destination for .md file
                    md_filename = item.stem + '.md'
                    if dest_category:
                        md_dest = self.config['dest_root'] / dest_category / md_filename
                    else:
                        md_dest = self.config['dest_root'] / '_unsorted' / md_filename

                    # Ensure directories exist
                    md_dest.parent.mkdir(parents=True, exist_ok=True)
                    scripts_backup.mkdir(parents=True, exist_ok=True)

                    # Handle filename conflicts for .md
                    final_md_dest = md_dest
                    if final_md_dest.exists():
                        counter = 1
                        while final_md_dest.exists():
                            final_md_dest = md_dest.parent / f"{item.stem}_{counter}.md"
                            counter += 1

                    # Write .md version
                    final_md_dest.write_text(md_content, encoding='utf-8')

                    # Move original to doc_scripts/
                    backup_dest = scripts_backup / item.name
                    if backup_dest.exists():
                        counter = 1
                        base = item.stem
                        while backup_dest.exists():
                            backup_dest = scripts_backup / f"{base}_{counter}{item.suffix}"
                            counter += 1

                    shutil.move(str(item), str(backup_dest))

                    # Log success
                    self.script_log.append({
                        'source_path': str(item),
                        'filename': item.name,
                        'file_type': ext_lower,
                        'md_destination': str(final_md_dest),
                        'backup_location': str(backup_dest),
                        'category': dest_category or '_unsorted',
                        'status': 'SUCCESS',
                    })
                    self.script_stats['converted'] += 1
                    print(f"    [SCRIPT] {item.name} -> {dest_category or '_unsorted'}/{md_filename}")

                except Exception as e:
                    self.script_log.append({
                        'source_path': str(item),
                        'filename': item.name,
                        'file_type': ext_lower,
                        'md_destination': '',
                        'backup_location': '',
                        'category': '',
                        'status': f'ERROR: {e}',
                    })
                    self.script_stats['errors'] += 1
                    print(f"    [ERROR] {item.name}: {e}")

    def _create_script_markdown(self, filename: str, content: str, lang: str) -> str:
        """Create markdown file with script content in code block."""
        # Extract a title from filename
        title = filename.replace('_', ' ').replace('-', ' ')
        title = re.sub(r'\.[^.]+$', '', title)  # Remove extension
        title = title.title()

        # Build markdown content
        lines = [
            f"# {title}",
            "",
            f"**Original file:** `{filename}`",
            f"**Type:** {lang.upper()}",
            "",
            "---",
            "",
            f"```{lang}",
            content,
            "```",
            "",
        ]
        return '\n'.join(lines)

    def _find_script_category(self, script_path: Path, content: str) -> str:
        """Try to find the best category for a script file."""
        filename = script_path.stem.lower()
        source_folder = script_path.parent.name.lower()

        # Extract keywords from filename
        filename_keywords = set(re.split(r'[_\-\s]+', filename))

        # Extract keywords from content (first 500 chars)
        content_sample = content[:500].lower()
        content_keywords = set(re.findall(r'\b[a-zA-Z_]{3,}\b', content_sample))

        # Try to match against categories
        best_match = None
        best_score = 0

        for cat_name, cat_info in self.parser.categories.items():
            score = 0

            # Check if source folder matches category
            if source_folder in str(cat_info['path']).lower():
                score += 0.4

            # Check filename keywords
            include_kw = set(cat_info['include_keywords'])
            fn_matches = filename_keywords & include_kw
            if fn_matches:
                score += len(fn_matches) * 0.15

            # Check content keywords
            content_matches = content_keywords & include_kw
            if content_matches:
                score += min(len(content_matches) * 0.05, 0.2)

            if score > best_score:
                best_score = score
                best_match = cat_name

        return best_match if best_score >= 0.15 else None

    def _process_docx(self):
        """Convert Word documents (.docx) to markdown."""
        if not self.config.get('docx_enabled', False):
            print("    [SKIP] Word document conversion disabled")
            return

        if not DOCX_AVAILABLE:
            print("    [SKIP] python-docx not installed (pip install python-docx)")
            return

        scripts_backup = self.config.get('scripts_backup_dir')

        for source_dir in self.config['source_dirs']:
            if not source_dir.exists():
                continue

            for item in self._get_files(source_dir, '*.docx'):
                if not item.is_file():
                    continue

                # Skip temp files (like ~$*.docx)
                if item.name.startswith('~$'):
                    self.docx_stats['skipped'] += 1
                    continue

                # Skip files matching skip patterns
                if self._should_skip(item):
                    continue

                try:
                    # Read Word document
                    doc = DocxDocument(str(item))

                    # Convert to markdown
                    md_content = self._docx_to_markdown(item.name, doc)
                    text_content = '\n'.join([p.text for p in doc.paragraphs])

                    # Find category
                    dest_category = self._find_script_category(item, text_content)

                    # Determine destination
                    md_filename = item.stem + '.md'
                    if dest_category:
                        md_dest = self.config['dest_root'] / dest_category / md_filename
                    else:
                        md_dest = self.config['dest_root'] / '_unsorted' / md_filename

                    # Ensure directories exist
                    md_dest.parent.mkdir(parents=True, exist_ok=True)
                    if scripts_backup:
                        scripts_backup.mkdir(parents=True, exist_ok=True)

                    # Handle filename conflicts
                    final_md_dest = md_dest
                    if final_md_dest.exists():
                        counter = 1
                        while final_md_dest.exists():
                            final_md_dest = md_dest.parent / f"{item.stem}_{counter}.md"
                            counter += 1

                    # Write markdown
                    final_md_dest.write_text(md_content, encoding='utf-8')

                    # Move original to doc_scripts/
                    backup_dest = None
                    if scripts_backup:
                        backup_dest = scripts_backup / item.name
                        if backup_dest.exists():
                            counter = 1
                            base = item.stem
                            while backup_dest.exists():
                                backup_dest = scripts_backup / f"{base}_{counter}.docx"
                                counter += 1
                        shutil.move(str(item), str(backup_dest))

                    # Log success
                    self.docx_log.append({
                        'source_path': str(item),
                        'filename': item.name,
                        'md_destination': str(final_md_dest),
                        'backup_location': str(backup_dest) if backup_dest else '',
                        'category': dest_category or '_unsorted',
                        'status': 'SUCCESS',
                    })
                    self.docx_stats['converted'] += 1
                    print(f"    [DOCX] {item.name} -> {dest_category or '_unsorted'}/{md_filename}")

                except Exception as e:
                    self.docx_log.append({
                        'source_path': str(item),
                        'filename': item.name,
                        'md_destination': '',
                        'backup_location': '',
                        'category': '',
                        'status': f'ERROR: {e}',
                    })
                    self.docx_stats['errors'] += 1
                    print(f"    [ERROR] {item.name}: {e}")

    def _docx_to_markdown(self, filename: str, doc) -> str:
        """Convert Word document to markdown format."""
        # Title from filename
        title = filename.replace('_', ' ').replace('-', ' ')
        title = re.sub(r'\.docx$', '', title, flags=re.IGNORECASE)
        title = title.title()

        lines = [
            f"# {title}",
            "",
            f"**Original file:** `{filename}`",
            f"**Type:** Word Document",
            "",
            "---",
            "",
        ]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            # Check paragraph style for headings
            style_name = para.style.name.lower() if para.style else ''

            if 'heading 1' in style_name:
                lines.append(f"## {text}")
            elif 'heading 2' in style_name:
                lines.append(f"### {text}")
            elif 'heading 3' in style_name:
                lines.append(f"#### {text}")
            elif 'title' in style_name:
                lines.append(f"## {text}")
            elif 'list' in style_name or text.startswith(('-', '*', '•')):
                # Clean up bullet points
                clean_text = text.lstrip('-*• \t')
                lines.append(f"- {clean_text}")
            else:
                lines.append(text)

            lines.append("")

        return '\n'.join(lines)

    def _process_html_files(self):
        """Move HTML files to appropriate category folders (as-is, no conversion)."""
        html_extensions = set(self.config.get('html_extensions', ['.html', '.htm']))

        for source_dir in self.config['source_dirs']:
            if not source_dir.exists():
                continue

            for item in self._get_files(source_dir):
                if not item.is_file():
                    continue

                ext_lower = item.suffix.lower()
                if ext_lower not in html_extensions:
                    continue

                # Skip files matching skip patterns
                if self._should_skip(item):
                    continue

                try:
                    # Read file content for scoring (just first part for keywords)
                    try:
                        content = item.read_text(encoding='utf-8', errors='ignore')[:5000]
                    except:
                        content = ""

                    # Score the file to find best category
                    filename = item.stem  # filename without extension
                    source_folder = item.parent.name

                    # Build keywords from filename, folder, and content
                    keywords = set()
                    keywords.update(re.split(r'[_\-\s]+', filename.lower()))
                    keywords.update(re.split(r'[_\-\s]+', source_folder.lower()))

                    # Extract some keywords from content (title, headings)
                    title_match = re.search(r'<title[^>]*>([^<]+)</title>', content, re.IGNORECASE)
                    if title_match:
                        keywords.update(re.split(r'[_\-\s]+', title_match.group(1).lower()))

                    h1_matches = re.findall(r'<h1[^>]*>([^<]+)</h1>', content, re.IGNORECASE)
                    for h1 in h1_matches[:3]:
                        keywords.update(re.split(r'[_\-\s]+', h1.lower()))

                    # Score against categories
                    best_category = None
                    best_score = 0

                    for cat_name, cat_info in self.scorer.categories.items():
                        cat_keywords = set(cat_info.get('include_keywords', []))
                        score = len(keywords & cat_keywords)
                        if score > best_score:
                            best_score = score
                            best_category = cat_name

                    # Determine destination
                    if best_category and best_score >= 1:
                        dest_category = best_category
                    else:
                        dest_category = '_unsorted'

                    # Build destination path
                    dest_dir = self.config['dest_root'] / dest_category
                    dest_dir.mkdir(parents=True, exist_ok=True)

                    dest_path = dest_dir / item.name

                    # Handle filename conflicts
                    final_dest = dest_path
                    if final_dest.exists():
                        counter = 1
                        while final_dest.exists():
                            final_dest = dest_dir / f"{item.stem}_{counter}{item.suffix}"
                            counter += 1

                    # MOVE the file (not copy)
                    shutil.move(str(item), str(final_dest))

                    # Log success
                    self.html_log.append({
                        'source_path': str(item),
                        'filename': item.name,
                        'destination': str(final_dest),
                        'category': dest_category,
                        'score': best_score,
                        'status': 'SUCCESS',
                    })
                    self.html_stats['moved'] += 1
                    print(f"    [HTML] {item.name} -> {dest_category}/")

                except Exception as e:
                    self.html_log.append({
                        'source_path': str(item),
                        'filename': item.name,
                        'destination': '',
                        'category': '',
                        'score': 0,
                        'status': f'ERROR: {e}',
                    })
                    self.html_stats['errors'] += 1
                    print(f"    [ERROR] {item.name}: {e}")

    def _process_old_tools(self):
        """Move old tool files (diagnostic scripts, batch files) to cleaning_up_old_script_tools/."""
        old_tools_dir = self.config.get('old_tools_dir')
        tool_patterns = self.config.get('old_tool_patterns', [])

        if not old_tools_dir:
            print("    [SKIP] Old tools directory not configured")
            return

        if not tool_patterns:
            print("    [SKIP] No old tool patterns configured")
            return

        # Ensure destination exists
        old_tools_dir.mkdir(parents=True, exist_ok=True)

        for source_dir in self.config['source_dirs']:
            if not source_dir.exists():
                continue

            for item in self._get_files(source_dir):
                if not item.is_file():
                    continue

                # Skip files matching skip patterns (module files, etc.)
                if self._should_skip(item):
                    continue

                # Check if file matches any old tool pattern
                is_tool = False
                matched_pattern = None
                for pattern in tool_patterns:
                    if re.match(pattern, item.name, re.IGNORECASE):
                        is_tool = True
                        matched_pattern = pattern
                        break

                if not is_tool:
                    continue

                try:
                    # Build destination path
                    dest_path = old_tools_dir / item.name

                    # Handle filename conflicts
                    final_dest = dest_path
                    if final_dest.exists():
                        counter = 1
                        while final_dest.exists():
                            final_dest = old_tools_dir / f"{item.stem}_{counter}{item.suffix}"
                            counter += 1

                    # Move the file
                    shutil.move(str(item), str(final_dest))

                    # Log success
                    self.old_tools_log.append({
                        'source_path': str(item),
                        'filename': item.name,
                        'file_type': item.suffix,
                        'destination': str(final_dest),
                        'matched_pattern': matched_pattern,
                        'status': 'SUCCESS',
                    })
                    self.old_tools_stats['moved'] += 1
                    print(f"    [OLD TOOL] {item.name} -> cleaning_up_old_script_tools/")

                except Exception as e:
                    self.old_tools_log.append({
                        'source_path': str(item),
                        'filename': item.name,
                        'file_type': item.suffix,
                        'destination': '',
                        'matched_pattern': matched_pattern,
                        'status': f'ERROR: {e}',
                    })
                    self.old_tools_stats['errors'] += 1
                    print(f"    [ERROR] {item.name}: {e}")

    def _is_old_tool(self, file_path: Path) -> bool:
        """Check if a file matches old tool patterns."""
        tool_patterns = self.config.get('old_tool_patterns', [])
        for pattern in tool_patterns:
            if re.match(pattern, file_path.name, re.IGNORECASE):
                return True
        return False

    def _find_asset_category(self, asset_path: Path) -> str:
        """Try to find the best category for an asset file."""
        # Get asset info
        filename = asset_path.stem.lower()
        source_folder = asset_path.parent.name.lower()

        # Extract keywords from filename
        filename_keywords = set(re.split(r'[_\-\s]+', filename))

        # Try to match against categories
        best_match = None
        best_score = 0

        for cat_name, cat_info in self.parser.categories.items():
            score = 0

            # Check if source folder matches category path
            if source_folder in str(cat_info['path']).lower():
                score += 0.5

            # Check filename keywords against category keywords
            include_kw = set(cat_info['include_keywords'])
            matches = filename_keywords & include_kw
            if matches:
                score += len(matches) * 0.2

            if score > best_score:
                best_score = score
                best_match = cat_name

        # Only return if we have reasonable confidence
        return best_match if best_score >= 0.2 else None

    def _find_remaining_content(self) -> list:
        """Find files not processed by any handler."""
        remaining = []
        asset_extensions = set(self.config.get('asset_extensions', []))
        script_extensions = set(self.config.get('script_extensions', {}).keys())

        for source_dir in self.config['source_dirs']:
            if not source_dir.exists():
                continue

            for item in self._get_files(source_dir):
                if not item.is_file():
                    continue
                # Skip .md files (already processed)
                if item.suffix.lower() == '.md':
                    continue
                # Skip asset files (already processed)
                if item.suffix.lower() in asset_extensions:
                    continue
                # Skip script files (already processed)
                if item.suffix.lower() in script_extensions:
                    continue
                # Skip .docx files (already processed)
                if item.suffix.lower() == '.docx':
                    continue
                # Skip HTML files (already processed)
                html_extensions = set(self.config.get('html_extensions', ['.html', '.htm']))
                if item.suffix.lower() in html_extensions:
                    continue
                # Skip old tool files (already processed)
                if self._is_old_tool(item):
                    continue

                remaining.append({
                    'path': str(item),
                    'name': item.name,
                    'type': item.suffix or 'no extension',
                    'folder': item.parent.name,
                })

        return remaining

    def _cleanup_empty_folders(self) -> int:
        """Remove empty folders from source directories. Returns count of removed folders."""
        removed_count = 0

        for source_dir in self.config['source_dirs']:
            if not source_dir.exists():
                continue

            # Walk bottom-up so we can remove nested empty folders
            for dirpath, dirnames, filenames in os.walk(str(source_dir), topdown=False):
                dir_path = Path(dirpath)

                # Skip if it's the source root itself
                if dir_path == source_dir:
                    continue

                # Check if directory is empty (no files and no subdirs)
                try:
                    contents = list(dir_path.iterdir())
                    if not contents:
                        dir_path.rmdir()
                        removed_count += 1
                        print(f"    [REMOVED] Empty folder: {dir_path.name}")
                except Exception as e:
                    print(f"    [ERROR] Could not remove {dir_path}: {e}")

        # Also try to remove the source dir itself if empty
        for source_dir in self.config['source_dirs']:
            if source_dir.exists():
                try:
                    contents = list(source_dir.iterdir())
                    if not contents:
                        source_dir.rmdir()
                        removed_count += 1
                        print(f"    [REMOVED] Empty source folder: {source_dir.name}")
                except:
                    pass

        return removed_count

    def _process_directory(self, source_dir: Path):
        """Process all .md files in a directory."""
        for md_file in self._get_files(source_dir, '*.md'):
            # Skip based on patterns
            if self._should_skip(md_file):
                continue

            # Score the file
            scores = self.scorer.score_file(md_file)
            best_match = scores[0] if scores else (self.config['unmatched_dest'], 0, ['no_match'])

            dest_folder = best_match[0]
            score = best_match[1]
            reasons = best_match[2]

            # Determine confidence
            confidence = self._get_confidence(score)

            # Only VERY_LOW (<15%) goes to _unsorted - everything else goes to matched folder
            if confidence == 'VERY_LOW':
                dest_folder = self.config['unmatched_dest']

            # Build destination path
            dest_path = self.config['dest_root'] / dest_folder / md_file.name

            # Execute move
            status, final_dest = self._move_file(md_file, dest_path)

            # Log the move
            self.move_log.append({
                'source_path': str(md_file),
                'filename': md_file.name,
                'source_folder': md_file.parent.name,
                'destination': str(final_dest) if final_dest else dest_folder,
                'confidence': confidence,
                'score': f"{score:.2f}",
                'match_reason': '; '.join(reasons),
                'status': status,
                'needs_review': 'YES' if confidence in ('LOW', 'MEDIUM', 'VERY_LOW') else '',
            })

            # Update stats
            if status == 'SUCCESS':
                self.stats['moved'] += 1
                self.stats[confidence.lower()] += 1
            elif status == 'SKIPPED':
                self.stats['skipped'] += 1
            else:
                self.stats['errors'] += 1

    def _should_skip(self, file_path: Path) -> bool:
        """Check if file should be skipped."""
        filename = file_path.name
        for pattern in self.config['skip_patterns']:
            if re.match(pattern, filename):
                return True
        return False

    def _get_confidence(self, score: float) -> str:
        """Convert score to confidence level."""
        if score >= self.config['high_confidence']:
            return 'HIGH'
        elif score >= self.config['medium_confidence']:
            return 'MEDIUM'
        elif score >= self.config['very_low_confidence']:
            return 'LOW'
        else:
            return 'VERY_LOW'

    def _move_file(self, source: Path, dest: Path) -> tuple:
        """
        Move a single file.
        Returns (status, final_dest_path)
        """
        try:
            if not source.exists():
                return ('ERROR: Source not found', None)

            # Ensure destination directory exists
            dest.parent.mkdir(parents=True, exist_ok=True)

            # Handle filename conflicts
            final_dest = dest
            if final_dest.exists():
                base = dest.stem
                ext = dest.suffix
                counter = 1
                while final_dest.exists():
                    final_dest = dest.parent / f"{base}_{counter}{ext}"
                    counter += 1

            # Copy file (preserving metadata)
            shutil.copy2(source, final_dest)

            # DELETE ORIGINAL - We're committing to the move
            source.unlink()

            return ('SUCCESS', final_dest)

        except Exception as e:
            return (f'ERROR: {e}', None)

    def _write_report(self) -> Path:
        """Write move log to CSV (separate files for docs and assets)."""
        timestamp = datetime.now().strftime('%Y-%m-%d_%H%M%S')

        # Ensure directory exists
        self.config['reports_dir'].mkdir(parents=True, exist_ok=True)

        # Write docs report
        docs_report = self.config['reports_dir'] / f'consolidation_docs_{timestamp}.csv'
        fieldnames = [
            'source_path', 'filename', 'source_folder',
            'destination', 'confidence', 'score', 'match_reason',
            'status', 'needs_review'
        ]
        with open(docs_report, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(self.move_log)

        # Write assets report (if any assets were processed)
        if self.asset_log:
            assets_report = self.config['reports_dir'] / f'consolidation_assets_{timestamp}.csv'
            asset_fieldnames = ['source_path', 'filename', 'file_type', 'destination', 'category', 'status']
            with open(assets_report, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=asset_fieldnames)
                writer.writeheader()
                writer.writerows(self.asset_log)

        # Write scripts report (if any scripts were processed)
        if self.script_log:
            scripts_report = self.config['reports_dir'] / f'consolidation_scripts_{timestamp}.csv'
            script_fieldnames = ['source_path', 'filename', 'file_type', 'md_destination', 'backup_location', 'category', 'status']
            with open(scripts_report, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=script_fieldnames)
                writer.writeheader()
                writer.writerows(self.script_log)

        # Write docx report (if any Word docs were processed)
        if self.docx_log:
            docx_report = self.config['reports_dir'] / f'consolidation_docx_{timestamp}.csv'
            docx_fieldnames = ['source_path', 'filename', 'md_destination', 'backup_location', 'category', 'status']
            with open(docx_report, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=docx_fieldnames)
                writer.writeheader()
                writer.writerows(self.docx_log)

        # Write old tools report (if any old tools were processed)
        if self.old_tools_log:
            tools_report = self.config['reports_dir'] / f'consolidation_old_tools_{timestamp}.csv'
            tools_fieldnames = ['source_path', 'filename', 'file_type', 'destination', 'matched_pattern', 'status']
            with open(tools_report, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=tools_fieldnames)
                writer.writeheader()
                writer.writerows(self.old_tools_log)

        # Write HTML report (if any HTML files were processed)
        if self.html_log:
            html_report = self.config['reports_dir'] / f'consolidation_html_{timestamp}.csv'
            html_fieldnames = ['source_path', 'filename', 'destination', 'category', 'score', 'status']
            with open(html_report, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=html_fieldnames)
                writer.writeheader()
                writer.writerows(self.html_log)

        return docs_report

    def _print_summary(self, report_path: Path, empty_removed: int = 0, remaining_files: list = None):
        """Print summary of what happened."""
        remaining_files = remaining_files or []

        print(f"\n{'='*60}")
        print("CONSOLIDATION COMPLETE")
        print('='*60)
        print(f"\nMarkdown Files:")
        print(f"  Moved successfully:  {self.stats['moved']}")
        print(f"    - High confidence (70%+):    {self.stats['high']}")
        print(f"    - Medium confidence (30-70%): {self.stats['medium']} (review recommended)")
        print(f"    - Low confidence (15-30%):    {self.stats['low']} (review recommended)")
        print(f"    - Very low (<15%):            {self.stats['very_low']} (sent to _unsorted/)")
        print(f"  Skipped:             {self.stats['skipped']}")
        print(f"  Errors:              {self.stats['errors']}")

        print(f"\nAsset Files (images, diagrams):")
        print(f"  Moved to _assets/:   {self.asset_stats['moved']}")
        print(f"  Errors:              {self.asset_stats['errors']}")

        print(f"\nScript Files (sql, json, etc.):")
        print(f"  Converted to .md:    {self.script_stats['converted']}")
        print(f"  Originals backed up: {self.script_stats['converted']} (in doc_scripts/)")
        print(f"  Errors:              {self.script_stats['errors']}")

        print(f"\nHTML Files:")
        print(f"  Moved (as-is):       {self.html_stats['moved']}")
        print(f"  Errors:              {self.html_stats['errors']}")

        print(f"\nWord Documents (.docx):")
        print(f"  Converted to .md:    {self.docx_stats['converted']}")
        print(f"  Skipped (temp files): {self.docx_stats['skipped']}")
        print(f"  Errors:              {self.docx_stats['errors']}")

        print(f"\nOld Tool Files (.py, .bat, .sh):")
        print(f"  Moved to cleanup folder: {self.old_tools_stats['moved']}")
        print(f"  Errors:              {self.old_tools_stats['errors']}")

        print(f"\nCleanup:")
        print(f"  Empty folders removed: {empty_removed}")

        # Show remaining content
        if remaining_files:
            print(f"\nRemaining Content (non-.md files): {len(remaining_files)}")
            # Group by file type
            by_type = defaultdict(list)
            for f in remaining_files:
                by_type[f['type']].append(f['name'])

            for ftype, files in sorted(by_type.items()):
                print(f"  {ftype}: {len(files)} files")
                for fname in files[:5]:  # Show first 5
                    print(f"    - {fname}")
                if len(files) > 5:
                    print(f"    ... and {len(files) - 5} more")

            print(f"\n  [!] These files need manual review!")
        else:
            print(f"\nRemaining Content: None (source folder clean!)")

        print(f"\nReport saved to:")
        print(f"  {report_path}")
        print(f"\nNext steps:")
        print(f"  1. Open the CSV report")
        print(f"  2. Filter by 'needs_review' = YES")
        print(f"  3. Manually move any misplaced files")
        print(f"  4. Check _unsorted/ for very low confidence files")
        if remaining_files:
            print(f"  5. Review remaining non-.md content listed above")
        print('='*60)


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def print_help():
    """Print usage help."""
    print("""
Documentation Consolidator
==========================

Usage:
    python doc_consolidator.py run       Auto-move all files + generate report
    python doc_consolidator.py --help    Show this help

What it does:
    1. Scans source directories for .md files
    2. Scores each file against destination categories (using _README.md keywords)
    3. AUTO-MOVES all .md files to best-match destination
    4. AUTO-MOVES asset files (images) to _assets/ folders
    5. Cleans up empty folders
    6. Generates cleanup reports (CSV) for post-move review

Confidence levels (for .md files):
    HIGH (70%+)     - Moved to matched folder, no flag
    MEDIUM (30-70%) - Moved to matched folder, flagged for review
    LOW (15-30%)    - Moved to matched folder, flagged for review
    VERY LOW (<15%) - Moved to _unsorted/, flagged

Asset files (.png, .jpg, .svg, .webp, etc.):
    - Moved to _assets/ folder within matched category
    - Or root _assets/ if no category match

Reports saved to:
    {reports_dir}
    """.format(reports_dir=CONFIG['reports_dir']))


def cmd_run():
    """Run full consolidation."""
    consolidator = Consolidator(CONFIG)
    consolidator.run()


def main():
    """Main entry point."""
    if len(sys.argv) < 2 or sys.argv[1] in ('--help', '-h', 'help'):
        print_help()
        sys.exit(0)

    command = sys.argv[1].lower()

    if command == 'run':
        cmd_run()
    else:
        print(f"Unknown command: {command}")
        print_help()
        sys.exit(1)


if __name__ == '__main__':
    main()
